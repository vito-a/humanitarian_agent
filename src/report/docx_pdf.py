"""
docx_pdf.py — Rendering helpers for humanitarian MVP reports.

Inputs expected:
- sections: Dict[str, List[{"title": str|None, "summary": str, "key_points": str|None, "ref_num": int|None, "published_at": str|None, "url": str, "source_site": str|None, ...}]]
- references_by_num: Dict[int|str, {"url": str, "title": str|None, "as_of" or "published_at": str|None}]
- exec_overview: Optional[str] — 1–3 paragraphs of overview (can be empty)
- filename: Output file path ('.docx' or '.pdf')
- country: Lowercase key like "ukraine" (used in headings)
- report_title: Optional override for the document title
- section_order: Optional[List[str]] — explicit order of sections

This module uses:
- python-docx for .docx
- reportlab for .pdf
"""

import re
from typing import Dict, List, Any, Optional, Iterable
from pathlib import Path
from math import ceil

# --- DOCX ---
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

# --- PDF (ReportLab) ---
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image
from reportlab.lib.units import inch
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.pagesizes import A4, LETTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from xml.sax.saxutils import escape as _esc

from ..config import PDF_LEFT_MARGIN, PDF_RIGHT_MARGIN, PDF_TOP_MARGIN, PDF_BOTTOM_MARGIN, AGENT_VERSION, ARIAL_FONT_PATH
from ..utils.branding import branding_line
from ..utils.coat_of_arms import ensure_coat_of_arms

# ------------------------ Common helpers ------------------------

def _pretty_country(country: str) -> str:
    return (country or "").strip().title()

def _default_title(country: str) -> str:
    return f"Humanitarian Situation Digest: {_pretty_country(country)}"

def _iter_sections_in_order(sections: Dict[str, Any], order: Optional[Iterable[str]]):
    if order:
        for k in order:
            if k in sections:
                yield k, sections[k]
    else:
        for k, v in sections.items():
            yield k, v

def _sorted_ref_nums(refs: dict) -> list:
    """Sort reference keys numerically even if JSON loaded them as strings."""
    def _k(n):
        try:
            return int(n)
        except Exception:
            return n
    return sorted(refs.keys(), key=_k)

def _ymd(dt_str: Optional[str]) -> str:
    """Format ISO-like or RFC3339-like dates to YYYY-MM-DD; fallback empty."""
    if not dt_str:
        return ""
    # assume first 10 chars are date in most sources
    return str(dt_str)[:10]

_SENT_SPLIT = re.compile(r"(?<=[.!?])\s+")

def _shorten_key_points_twofold(text: str) -> str:
    """
    Reduce key points roughly two-fold by sentence count:
    keep the first ceil(n/2) sentences.
    """
    t = (text or "").strip()
    if not t:
        return ""
    sents = _SENT_SPLIT.split(t)
    if len(sents) <= 1:
        return t
    keep = max(1, ceil(len(sents) / 2))
    return " ".join(sents[:keep]).strip()

def add_coat_of_arms(doc, country: str):
    """
    Insert the coat of arms image left-aligned at the top of the report.
    """
    try:
        coa_path = ensure_coat_of_arms(country)
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        with coa_path.open("rb") as f:
            run.add_picture(f, width=Inches(0.9))
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # left alignment
    except Exception as e:
        print(f"⚠️ Could not add coat of arms image: {e}")

def coat_of_arms_image(country: str):
    """
    Return a flowable Image of the coat of arms, left-aligned at top.
    """
    try:
        coa_path = ensure_coat_of_arms(country)
        img = Image(coa_path, width=0.9*inch, height=0.9*inch)
        img.hAlign = "LEFT"   # left alignment
        return img
    except Exception as e:
        print(f"⚠️ Could not load coat of arms image: {e}")
        return None

def add_header_with_coa_docx(doc, country: str, report_title: str,
                             img_width_in: float = 0.9, gap_in: float = 0.2):
    """
    Insert a header row:
      [ coat of arms image ][ gap ][ report title ]
    - Table spans the full page width.
    - Image and gap columns fixed; title column takes remaining width.
    - Vertically centered, no borders.
    """
    # Create table
    #tbl = doc.add_table(rows=1, cols=3)
    tbl = doc.add_table(rows=1, cols=2)
    _strip_table_borders(tbl)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Compute usable page width (page minus left+right margins)
    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    # Convert to EMUs (1 inch = 914400 EMUs)
    emu_per_inch = 914400
    img_width_emu = int(img_width_in * emu_per_inch)
    gap_width_emu = int(gap_in * emu_per_inch)
    title_width_emu = usable_width * emu_per_inch - img_width_emu - gap_width_emu

    tbl.columns[0].width = img_width_emu
    #tbl.columns[1].width = gap_width_emu
    #tbl.columns[2].width = title_width_emu
    tbl.columns[1].width = title_width_emu

    row = tbl.rows[0]
    #cell_img, cell_gap, cell_title = row.cells
    cell_img, cell_title = row.cells

    # Image cell
    p_img = cell_img.paragraphs[0]
    run = p_img.add_run()
    try:
        coa_path = ensure_coat_of_arms(country)
        with coa_path.open("rb") as f:
            run.add_picture(f, width=Inches(0.9))
    except Exception:
        p_img.add_run(" ")
    p_img.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell_img.width = img_width_in
    cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Gap cell
    #p_gap = cell_gap.paragraphs[0]
    #p_gap.add_run(" ")
    #cell_gap.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Title cell
    p_title = cell_title.paragraphs[0]
    run_t = p_title.add_run(report_title)
    if "Title" in doc.styles:
        p_title.style = doc.styles["Title"]
    else:
        run_t.font.size = Pt(22)
        run_t.font.bold = True
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell_title.width = usable_width - img_width_in
    cell_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def header_with_coa_pdf(country: str, report_title: str, styles,
                        img_size_in: float = 0.9, gap_in: float = 0.2,
                        total_width: float | None = None):
    """
    Returns a flowable that renders:
      [ coat of arms image ] [ gap ] [ report title ]
    left-aligned, vertically centered, no borders.
    - `styles` should include a "Title" style (e.g., getSampleStyleSheet()["Title"])
    - If `total_width` is provided (doc.width), we set explicit column widths.
    """
    coa_path = ensure_coat_of_arms(country)
    try:
        img = Image(coa_path, width=img_size_in*inch, height=img_size_in*inch)
    except Exception:
        # Fallback placeholder if image missing
        img = Paragraph("⚠️ Coat of arms not available", styles["Normal"])

    # Build the row: [img][gap][title]
    title_para = Paragraph(report_title, styles["Title"])
    data = [[img, Spacer(gap_in*inch, 0), title_para]]

    if total_width:
        img_w = img_size_in * inch
        gap_w = gap_in * inch
        title_w = total_width - img_w - gap_w - PDF_LEFT_MARGIN * inch - PDF_RIGHT_MARGIN * inch
        col_widths = [img_w, gap_w, title_w]
    else:
        col_widths = None  # let platypus auto-size

    tbl = Table(data, colWidths=col_widths, hAlign="LEFT")
    tbl.setStyle(TableStyle([
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
        ("LEFTPADDING", (0,0), (-1,-1), 0),
        ("RIGHTPADDING", (0,0), (-1,-1), 0),
        ("TOPPADDING", (0,0), (-1,-1), 0),
        ("BOTTOMPADDING", (0,0), (-1,-1), 2),
        # no borders:
        ("BOX", (0,0), (-1,-1), 0, colors.white),
        ("INNERGRID", (0,0), (-1,-1), 0, colors.white),
    ]))
    return tbl

def add_docx_header_branding(doc) -> None:
    """
    Add top-right small branding line in the header of ALL sections/pages.
    """
    text = branding_line(AGENT_VERSION)
    for section in doc.sections:
        header = section.header
        # ensure at least one paragraph exists
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # clear existing runs in that paragraph (optional)
        for r in list(para.runs):
            r.clear()
        run = para.add_run(text)
        run.font.size = Pt(8)
        run.font.italic = True   # optional for subtle styling

# docx: bottom-right branding in footer for every section/page
def add_docx_footer_branding(doc) -> None:
    """
    Add bottom-right small branding line in the FOOTER of ALL sections/pages.
    """
    text = branding_line(AGENT_VERSION)
    for section in doc.sections:
        footer = section.footer
        para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # clear existing runs (optional)
        for r in list(para.runs):
            r.clear()
        run = para.add_run(text)
        run.font.size = Pt(8)
        run.font.italic = True  # optional

def _on_page_branding_top(canvas: Canvas, doc) -> None:
    """
    Draw small, right-aligned branding text at the top of each page.
    Uses the DocTemplate margins to position.
    """
    text = branding_line(AGENT_VERSION)
    canvas.saveState()
    try:
        canvas.setFont("Helvetica", 8)
        # Right edge inside page margin
        right_x = doc.pagesize[0] - doc.rightMargin
        # Slightly below the top margin to avoid clipping
        y = doc.pagesize[1] - doc.topMargin + (0.22 * inch)
        # Draw right-aligned
        canvas.drawRightString(right_x, y, text)
    finally:
        canvas.restoreState()

def _on_page_branding_bottom(canvas: Canvas, doc) -> None:
    """
    Draw small, right-aligned branding text at the bottom of each page.
    """
    text = branding_line(AGENT_VERSION)  # already timezone-aware
    canvas.saveState()
    try:
        canvas.setFont("Helvetica", 8)
        right_x = doc.pagesize[0] - doc.rightMargin
        # Slightly above the bottom margin to avoid clipping
        y = doc.bottomMargin - (0.18 * inch)
        # If your bottom margin is small, clamp to a minimum
        if y < 0.25 * inch:
            y = 0.25 * inch
        canvas.drawRightString(right_x, y, text)
    finally:
        canvas.restoreState()

# ------------------------ DOCX utilities ------------------------

def _add_hyperlink(paragraph, url: str, text: str, color="0000FF", underline=True):
    """Create a clickable hyperlink within a paragraph with blue/underline style."""
    if not url:
        r = paragraph.add_run(text)
        return r
    # relationship id
    r_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    # hyperlink tag
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    # run with properties
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    if color:
        c = OxmlElement('w:color'); c.set(qn('w:val'), color); rPr.append(c)
    if underline:
        u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def _strip_table_borders(tbl):
    """Remove all table borders in a python-docx table."""
    tblPr = tbl._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    # Create/replace tblBorders with nil edges
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)

    # clear previous children
    for child in list(borders):
        borders.remove(child)

    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'), 'nil')
        borders.append(tag)

def _set_cell_bottom_border(cell, sz=12, color="000000"):
    """Add only a bottom border to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is None:
        borders = OxmlElement('w:tcBorders')
        tcPr.append(borders)
    # remove previous bottom if exists
    for child in list(borders):
        if child.tag == qn('w:bottom'):
            borders.remove(child)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(sz)))
    bottom.set(qn('w:color'), color)
    borders.append(bottom)

# ------------------ DOCX ------------------

def generate_docx_structured(
    sections: Dict[str, List[Dict[str, Any]]],
    references_by_num: Dict[int, Dict[str, Any]],
    exec_overview: str,
    filename: str,
    country: str,
    report_title: Optional[str] = None,
    section_order: Optional[List[str]] = None,
) -> None:
    doc = Document()

    # Add header with branding
    add_docx_footer_branding(doc)

    # Add header with title and the coat of arms
    title = report_title or _default_title(country)
    #h = doc.add_heading(title, 0)
    #h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    #add_coat_of_arms(doc, country)
    add_header_with_coa_docx(doc, country, title)

    # Executive summary (optional)
    if exec_overview:
        doc.add_heading("Executive Summary", level=1)
        for block in [b.strip() for b in exec_overview.split("\n\n") if b.strip()]:
            p = doc.add_paragraph()
            run = p.add_run(block)
            run.font.size = Pt(11)

    # Sections in provided order
    first_section = True
    for section_name, items in _iter_sections_in_order(sections, section_order):
        if not items:
            continue
        if first_section:
            first_section = False
        else:
            doc.add_page_break()   # page break before each section
            doc.add_heading(section_name, level=1)

        # Separate narrative (no ref_num) and article items
        article_items = [it for it in items if it.get("ref_num") is not None]
        narrative_items = [it for it in items if it.get("ref_num") is None]

        # Narrative blocks (Summary / Why support / Conclusion)
        for it in narrative_items:
            text = (it.get("summary") or "").strip()
            if text:
                for block in [b.strip() for b in text.split("\n\n") if b.strip()]:
                    p = doc.add_paragraph()
                    run = p.add_run(block)
                    run.font.size = Pt(11)

        # Articles: one Evidence—Key points table per section
        if article_items:
            tbl = doc.add_table(rows=len(article_items) + 1, cols=2)
            # Remove all borders and then add a single line below the header row
            _strip_table_borders(tbl)

            # Header row with bottom border only
            hdr0 = tbl.cell(0, 0).paragraphs[0].add_run("Evidence"); hdr0.bold = True
            hdr1 = tbl.cell(0, 1).paragraphs[0].add_run("Key points"); hdr1.bold = True
            for cell in tbl.rows[0].cells:
                _set_cell_bottom_border(cell, sz=16, color="000000")

            for r, it in enumerate(article_items, start=1):
                url = it.get("url", "")
                title_txt = it.get("short_title") or it.get("title") or "Untitled"
                src = it.get("source_site") or "Source"
                ref = it.get("ref_num")

                # Pick a date: prefer item's own; fallback to references_by_num
                dt_raw = it.get("published_at") or (references_by_num.get(ref, {}) if ref else {}).get("as_of") or \
                         (references_by_num.get(ref, {}) if ref else {}).get("published_at")
                date_ymd = _ymd(dt_raw)

                # Evidence cell: Title (hyperlink), Source (hyperlink), Date
                p_e = tbl.cell(r, 0).paragraphs[0]
                _add_hyperlink(p_e, url, title_txt, color="0000FF", underline=True)
                p_e2 = tbl.cell(r, 0).add_paragraph("— ")
                _add_hyperlink(p_e2, url, src, color="0000FF", underline=True)
                if date_ymd:
                    tbl.cell(r, 0).add_paragraph(date_ymd)

                # Key points cell (shortened by 2x; append [n])
                base_kp = (it.get("key_points") or it.get("summary") or "").strip()
                kp_short = base_kp # _shorten_key_points_twofold(base_kp)
                if ref and not kp_short.rstrip().endswith(f"[{ref}]"):
                    kp_short = f"{kp_short.rstrip()} [{ref}]"
                tbl.cell(r, 1).paragraphs[0].add_run(kp_short)

            for row in tbl.rows:
                for cell in row.cells:
                    _set_cell_bottom_border(cell, sz=8, color="000000")

    # References (with clickable [n], blue + underline, short dates, and full URL hyperlink visible)
    if references_by_num:
        doc.add_page_break()
        doc.add_heading("References", level=1)
        for num in _sorted_ref_nums(references_by_num):
            meta = references_by_num[num] or {}
            url = meta.get("url", "") or ""
            title_txt = meta.get("title") or url
            dt_short = _ymd(meta.get("as_of") or meta.get("published_at") or "")

            p = doc.add_paragraph()
            _add_hyperlink(p, url, f"[{num}]", color="0000FF", underline=True)
            p.add_run(" " + title_txt)
            if dt_short:
                p.add_run(f" — {dt_short}")
            if url:
                p.add_run(" — ")
                _add_hyperlink(p, url, url, color="0000FF", underline=True)

    Path(filename).parent.mkdir(parents=True, exist_ok=True)
    doc.save(filename)


# ------------------ PDF ------------------

def generate_pdf_structured(
    sections: Dict[str, List[Dict[str, Any]]],
    references_by_num: Dict[int, Dict[str, Any]],
    exec_overview: str,
    filename: str,
    country: str,
    report_title: Optional[str] = None,
    section_order: Optional[List[str]] = None,
) -> None:
#    pdfmetrics.registerFont(TTFont('Arial', ARIAL_FONT_PATH))
    title = report_title or _default_title(country)
    Path(filename).parent.mkdir(parents=True, exist_ok=True)

    # Styles
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle(name="H1", parent=styles["Heading1"], spaceBefore=6, spaceAfter=6)
    body = ParagraphStyle(name="Body", parent=styles["BodyText"], leading=14)

    flow = []
    #title_style = ParagraphStyle(name="TitleCenter", parent=styles["Title"], alignment=TA_CENTER)
    #title_style.fontName = 'Arial'
    #h1.fontName = 'Arial'
    #body.fontName = 'Arial'
    #flow.append(coat_of_arms_image(country))
    #flow.append(Spacer(1, 0.2*inch))
    #flow.append(Paragraph(title, title_style))
    #flow.append(Spacer(1, 12))
    flow.append(header_with_coa_pdf(country, title, styles, total_width=LETTER[0]))
    flow.append(Spacer(1, 0.2*inch))

    # Executive Summary
    if exec_overview:
        flow.append(Paragraph("Executive Summary", h1))
        flow.append(Spacer(1, 6))
        for block in [b.strip() for b in exec_overview.split("\n\n") if b.strip()]:
            flow.append(Paragraph(_esc(block), body))
            flow.append(Spacer(1, 6))
        flow.append(Spacer(1, 12))

    # Sections in order
    first_section = True
    for section_name, items in _iter_sections_in_order(sections, section_order):
        if not items:
            continue
        if first_section:
            first_section = False
        else:
            flow.append(PageBreak())   # page break before each section

        flow.append(Paragraph(_esc(section_name), h1))
        flow.append(Spacer(1, 6))

        # Narrative
        narrative_items = [it for it in items if it.get("ref_num") is None]
        for it in narrative_items:
            text = (it.get("summary") or "").strip()
            if text:
                for block in [b.strip() for b in text.split("\n\n") if b.strip()]:
                    flow.append(Paragraph(_esc(block), body))
                    flow.append(Spacer(1, 6))

        # Articles table
        article_items = [it for it in items if it.get("ref_num") is not None]
        if article_items:
            # Header
            data = [[Paragraph("<b>Evidence</b>", body), Paragraph("<b>Key points</b>", body)]]
            for it in article_items:
                url = it.get("url", "")
                title_txt = _esc(it.get("short_title") or it.get("title") or "Untitled")
                src = _esc(it.get("source_site") or "Source")
                ref = it.get("ref_num")

                # Date selection (same as DOCX) and format to YYYY-MM-DD
                meta = references_by_num.get(ref, {}) if ref else {}
                dt_raw = it.get("published_at") or meta.get("as_of") or meta.get("published_at")
                date_ymd = _ymd(dt_raw)

                # Evidence links (title + source, then date)
                ev_html = (
                    f'<link href="{url}" color="blue" underline="1">{title_txt}</link>'
                    f'<br/>— <link href="{url}" color="blue" underline="1">{src}</link>'
                )
                if date_ymd:
                    ev_html += f"<br/>{_esc(date_ymd)}"

                # Key points (shortened 2x; make [n] clickable, blue + underlined)
                base_kp = (it.get("key_points") or it.get("summary") or "").strip()
                kp_short = base_kp # _shorten_key_points_twofold(base_kp)
                if ref and f"[{ref}]" not in kp_short:
                    kp_short = f"{kp_short.rstrip()} [{ref}]"
                if ref:
                    kp_html = _esc(kp_short).replace(
                        f"[{ref}]",
                        f'<link href="{url}" color="blue" underline="1">[{ref}]</link>'
                    )
                else:
                    kp_html = _esc(kp_short)

                data.append([Paragraph(ev_html, body), Paragraph(kp_html, body)])

            tbl = Table(data, colWidths=[180, 300])

            # Clean style: only a line below the header row
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0,0), (-1,0), colors.whitesmoke),
                ("LINEBELOW", (0,0), (-1,-1), 0.5, colors.black),  # line under each row
                ("VALIGN", (0,0), (-1,-1), "TOP"),
                ("LEFTPADDING", (0,0), (-1,-1), 6),
                ("RIGHTPADDING", (0,0), (-1,-1), 6),
            ]))
            flow.append(tbl)
            flow.append(Spacer(1, 12))

    # References: [n] link, short date, and full URL as clickable text
    if references_by_num:
        flow.append(PageBreak())
        flow.append(Paragraph("References", h1))
        flow.append(Spacer(1, 6))
        for num in _sorted_ref_nums(references_by_num):
            meta = references_by_num[num] or {}
            url = meta.get("url", "") or ""
            title_txt = _esc(meta.get("title") or url)
            dt_short = _ymd(meta.get("as_of") or meta.get("published_at") or "")
            line = f'<link href="{url}" color="blue" underline="1">[{num}]</link> {title_txt}'
            if dt_short:
                line += f" — {_esc(dt_short)}"
            if url:
                line += f' — <link href="{url}" color="blue" underline="1">{_esc(url)}</link>'
            flow.append(Paragraph(line, body))
            flow.append(Spacer(1, 4))

    doc = SimpleDocTemplate(
        filename,
        pagesize = LETTER,
        title = title,
        leftMargin = PDF_LEFT_MARGIN * inch,
        rightMargin = PDF_RIGHT_MARGIN * inch,
        topMargin = PDF_TOP_MARGIN * inch,
        bottomMargin = PDF_BOTTOM_MARGIN * inch
    )
    doc.build(
        flow,
        onFirstPage=_on_page_branding_bottom,
        onLaterPages=_on_page_branding_bottom,
    )
